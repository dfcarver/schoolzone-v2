"""
Generate docs/SPECIFICATION.docx from docs/SPECIFICATION.md
Uses python-docx to convert markdown structure into a styled Word document.
"""

import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SPEC_MD = Path(__file__).parent / "SPECIFICATION.md"
OUT_DOCX = Path(__file__).parent / "SPECIFICATION.docx"

# ── Colour palette ────────────────────────────────────────────────────────────
NAVY   = RGBColor(0x1E, 0x3A, 0x5F)   # headings
TEAL   = RGBColor(0x0D, 0x7C, 0x8F)   # h2
SLATE  = RGBColor(0x44, 0x55, 0x66)   # h3
BLACK  = RGBColor(0x1A, 0x1A, 0x2E)
GREY   = RGBColor(0x55, 0x55, 0x55)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
TABLE_HEADER_BG = "1E3A5F"            # navy (hex string for XML)
TABLE_ALT_BG    = "F0F4F8"


def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_cell_borders(table):
    """Add subtle borders to every cell."""
    for row in table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            borders = OxmlElement("w:tcBorders")
            for side in ("top", "left", "bottom", "right"):
                el = OxmlElement(f"w:{side}")
                el.set(qn("w:val"), "single")
                el.set(qn("w:sz"), "4")
                el.set(qn("w:space"), "0")
                el.set(qn("w:color"), "CCCCCC")
                borders.append(el)
            tcPr.append(borders)


def apply_run_inline(run, text: str):
    """Apply bold/italic/code formatting to a run based on markers in text."""
    run.text = text


def add_paragraph_with_inline(doc, text: str, style=None):
    """Add a paragraph handling **bold**, *italic*, and `code` inline marks."""
    para = doc.add_paragraph(style=style)
    # Split on bold/italic/code markers
    pattern = r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)"
    parts = re.split(pattern, text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = para.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run = para.add_run(part[1:-1])
            run.italic = True
        elif part.startswith("`") and part.endswith("`"):
            run = para.add_run(part[1:-1])
            run.font.name = "Courier New"
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)
        else:
            para.add_run(part)
    return para


def style_heading(para, level: int):
    for run in para.runs:
        run.font.color.rgb = NAVY if level == 1 else (TEAL if level == 2 else SLATE)
        run.font.bold = True
        run.font.size = Pt(18 if level == 1 else (14 if level == 2 else 12))


def parse_and_build(doc: Document, lines: list[str]):
    i = 0
    in_code_block = False
    code_lines: list[str] = []
    in_table = False
    table_rows: list[list[str]] = []

    def flush_table():
        nonlocal in_table, table_rows
        if not table_rows:
            return
        # Filter out separator rows (---|--- etc.)
        data_rows = [r for r in table_rows if not all(re.match(r"^[-:]+$", c.strip()) for c in r)]
        if not data_rows:
            table_rows = []
            in_table = False
            return

        col_count = max(len(r) for r in data_rows)
        table = doc.add_table(rows=len(data_rows), cols=col_count)
        table.style = "Table Grid"
        set_cell_borders(table)

        for ri, row in enumerate(data_rows):
            for ci, cell_text in enumerate(row):
                if ci >= col_count:
                    break
                cell = table.cell(ri, ci)
                cell_text = cell_text.strip()
                # Strip inline formatting for table cells
                cell_text = re.sub(r"\*\*(.+?)\*\*", r"\1", cell_text)
                cell_text = re.sub(r"\*(.+?)\*", r"\1", cell_text)
                cell_text = re.sub(r"`(.+?)`", r"\1", cell_text)
                cell_text = re.sub(r"\[(.+?)\]\(.+?\)", r"\1", cell_text)  # strip links
                cell.text = cell_text
                para = cell.paragraphs[0]
                para.paragraph_format.space_before = Pt(2)
                para.paragraph_format.space_after = Pt(2)
                if ri == 0:
                    set_cell_bg(cell, TABLE_HEADER_BG)
                    for run in para.runs:
                        run.font.color.rgb = WHITE
                        run.font.bold = True
                        run.font.size = Pt(9)
                elif ri % 2 == 0:
                    set_cell_bg(cell, TABLE_ALT_BG)
                    for run in para.runs:
                        run.font.size = Pt(9)
                else:
                    for run in para.runs:
                        run.font.size = Pt(9)

        doc.add_paragraph()  # spacing after table
        table_rows = []
        in_table = False

    def flush_code(lines_buf):
        if not lines_buf:
            return
        para = doc.add_paragraph()
        para.paragraph_format.left_indent = Inches(0.4)
        para.paragraph_format.space_before = Pt(4)
        para.paragraph_format.space_after = Pt(4)
        run = para.add_run("\n".join(lines_buf))
        run.font.name = "Courier New"
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)
        # light grey background via shading
        pPr = para._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "F3F4F6")
        pPr.append(shd)

    while i < len(lines):
        line = lines[i]

        # ── Code blocks ───────────────────────────────────────────────────────
        if line.startswith("```"):
            if not in_code_block:
                if in_table:
                    flush_table()
                in_code_block = True
                code_lines = []
            else:
                flush_code(code_lines)
                in_code_block = False
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # ── Tables ────────────────────────────────────────────────────────────
        if line.startswith("|"):
            in_table = True
            cells = [c for c in line.split("|") if c != ""]
            table_rows.append(cells)
            i += 1
            continue
        else:
            if in_table:
                flush_table()

        # ── Horizontal rule ───────────────────────────────────────────────────
        if re.match(r"^---+$", line.strip()):
            para = doc.add_paragraph()
            pPr = para._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "6")
            bottom.set(qn("w:space"), "1")
            bottom.set(qn("w:color"), "CCCCCC")
            pBdr.append(bottom)
            pPr.append(pBdr)
            i += 1
            continue

        # ── Headings ──────────────────────────────────────────────────────────
        m = re.match(r"^(#{1,4})\s+(.*)", line)
        if m:
            level = len(m.group(1))
            text = m.group(2).strip()
            # Strip anchor suffixes like {#foo}
            text = re.sub(r"\s*\{#[^}]+\}", "", text)
            heading_style = f"Heading {level}"
            try:
                para = doc.add_heading(text, level=level)
            except Exception:
                para = doc.add_paragraph(text)
            style_heading(para, level)
            i += 1
            continue

        # ── Blockquote ────────────────────────────────────────────────────────
        if line.startswith(">"):
            text = line.lstrip("> ").strip()
            text = re.sub(r"\[(.+?)\]\(.+?\)", r"\1", text)
            para = doc.add_paragraph(style="Quote") if "Quote" in [s.name for s in doc.styles] else doc.add_paragraph()
            para.paragraph_format.left_indent = Inches(0.4)
            run = para.add_run(text)
            run.italic = True
            run.font.color.rgb = SLATE
            i += 1
            continue

        # ── Bullet list ───────────────────────────────────────────────────────
        m = re.match(r"^(\s*)[-*]\s+(.*)", line)
        if m:
            indent = len(m.group(1)) // 2
            text = m.group(2).strip()
            text = re.sub(r"\[(.+?)\]\(.+?\)", r"\1", text)
            para = add_paragraph_with_inline(doc, text, style="List Bullet")
            para.paragraph_format.left_indent = Inches(0.25 + indent * 0.25)
            para.paragraph_format.space_after = Pt(2)
            i += 1
            continue

        # ── Numbered list ─────────────────────────────────────────────────────
        m = re.match(r"^(\s*)\d+\.\s+(.*)", line)
        if m:
            text = m.group(2).strip()
            text = re.sub(r"\[(.+?)\]\(.+?\)", r"\1", text)
            para = add_paragraph_with_inline(doc, text, style="List Number")
            para.paragraph_format.space_after = Pt(2)
            i += 1
            continue

        # ── Blank line ────────────────────────────────────────────────────────
        if line.strip() == "":
            i += 1
            continue

        # ── Normal paragraph ──────────────────────────────────────────────────
        text = line.strip()
        text = re.sub(r"\[(.+?)\]\(.+?\)", r"\1", text)  # strip markdown links
        para = add_paragraph_with_inline(doc, text)
        para.paragraph_format.space_after = Pt(4)
        for run in para.runs:
            run.font.size = Pt(10)
            run.font.color.rgb = BLACK

        i += 1

    if in_table:
        flush_table()
    if in_code_block:
        flush_code(code_lines)


def build_docx():
    doc = Document()

    # ── Page margins ──────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3.0)
        section.right_margin  = Cm(2.5)

    # ── Default body font ─────────────────────────────────────────────────────
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10)
    doc.styles["Normal"].font.color.rgb = BLACK

    # ── Cover page ────────────────────────────────────────────────────────────
    doc.add_paragraph()
    doc.add_paragraph()
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run("SchoolZone Digital Twin")
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = NAVY

    sub_para = doc.add_paragraph()
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub_para.add_run("MRDT — Monitoring, Risk & Digital Twin")
    run.font.size = Pt(14)
    run.font.color.rgb = TEAL

    doc.add_paragraph()
    spec_para = doc.add_paragraph()
    spec_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = spec_para.add_run("Product Specification")
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = SLATE

    doc.add_paragraph()
    meta_para = doc.add_paragraph()
    meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta_para.add_run("Version 1.0  ·  Abu Dhabi Department of Transport  ·  Proprietary")
    run.font.size = Pt(10)
    run.font.color.rgb = GREY

    doc.add_page_break()

    # ── Parse markdown ────────────────────────────────────────────────────────
    md_text = SPEC_MD.read_text(encoding="utf-8")
    lines = md_text.splitlines()

    # Skip the front-matter title block (first 7 lines — heading + metadata)
    # We already have a cover page
    skip_until = 0
    for idx, l in enumerate(lines):
        if l.strip() == "---" and idx > 0:
            skip_until = idx + 1
            break

    parse_and_build(doc, lines[skip_until:])

    doc.save(OUT_DOCX)
    print(f"Saved: {OUT_DOCX}")


if __name__ == "__main__":
    build_docx()
